"""Тесты парсеров: фикстуры генерируются программно (pymupdf/python-docx) —
бинарники в репозитории не хранятся, вход детерминирован."""

import io

import pymupdf
import pytest
from docx import Document as DocxBuilder

from lyra.ingest.ir import BlockType
from lyra.ingest.parsers import detect_format, parse_document
from lyra.ingest.parsers.base import ParserError
from lyra.ingest.parsers.confluence_html import parse_confluence_html
from lyra.ingest.parsers.markdown import parse_markdown

# --- detect_format ---


def test_detect_format() -> None:
    assert detect_format(b"%PDF-1.7 ...", "a.pdf") == "pdf"
    assert detect_format(b"PK\x03\x04...", "a.docx") == "docx"
    assert detect_format(b"PK\x03\x04...", "a.zip") is None
    assert detect_format("# Заголовок".encode(), "readme.md") == "markdown"
    assert detect_format("просто текст".encode(), "note.txt") == "txt"
    assert detect_format(b"\xff\xfe\x00\x01\x02", "bin.dat") is None


# --- markdown ---

MD = """---
title: Политика отпусков
---
# Отпуск

Сотрудникам предоставляется 28 дней.

## Условия

| Стаж | Дней |
|------|------|
| 1 год | 28 |

```python
days = 28
```
"""


def test_markdown_structure() -> None:
    ir = parse_markdown(MD, title="fallback.md")
    assert ir.title == "Политика отпусков"  # из front-matter
    sections = ir.iter_sections()
    paths = [path for path, _ in sections]
    assert ["Отпуск"] in paths
    assert ["Отпуск", "Условия"] in paths
    blocks = [b for _, s in sections for b in s.blocks]
    assert any(b.type == BlockType.TABLE and "Стаж" in b.text for b in blocks)
    assert any(b.type == BlockType.CODE and b.meta.get("lang") == "python" for b in blocks)


def test_markdown_hash_deterministic() -> None:
    first = parse_markdown(MD, title="t").content_hash()
    second = parse_markdown(MD, title="t").content_hash()
    changed = parse_markdown(MD + "\nНовая строка.", title="t").content_hash()
    assert first == second
    assert first != changed


# --- docx ---


def _make_docx() -> bytes:
    builder = DocxBuilder()
    builder.add_heading("Инструкция", level=1)
    builder.add_paragraph("Первый шаг настройки.")
    builder.add_heading("Детали", level=2)
    builder.add_paragraph("Подробности процедуры.")
    table = builder.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    table.rows[1].cells[0].text = "Таймаут"
    table.rows[1].cells[1].text = "30с"
    buffer = io.BytesIO()
    builder.save(buffer)
    return buffer.getvalue()


def test_docx_structure() -> None:
    ir = parse_document(_make_docx(), fmt="docx", title="инструкция.docx")
    paths = [path for path, _ in ir.iter_sections()]
    assert ["Инструкция"] in paths
    assert ["Инструкция", "Детали"] in paths
    blocks = [b for _, s in ir.iter_sections() for b in s.blocks]
    tables = [b for b in blocks if b.type == BlockType.TABLE]
    assert tables and "Таймаут" in tables[0].text


def test_docx_broken_raises_parser_error() -> None:
    with pytest.raises(ParserError):
        parse_document(b"PK\x03\x04" + "мусор".encode(), fmt="docx", title="bad.docx")


# --- pdf ---


def _make_pdf(pages: int = 3) -> bytes:
    # Латиница: базовые PDF-шрифты (helv) не содержат кириллицу; парсер языконезависим
    document = pymupdf.open()
    for page_number in range(pages):
        page = document.new_page()
        page.insert_text((72, 60), f"Section {page_number + 1}", fontsize=18)
        for line in range(4):
            page.insert_text(
                (72, 100 + line * 20),
                f"Page {page_number + 1} content, line {line + 1} of the document.",
                fontsize=11,
            )
        page.insert_text((72, 800), "Confidential - Example LLC", fontsize=8)
    data: bytes = document.tobytes()
    document.close()
    return data


def test_pdf_headings_and_footer_removal() -> None:
    ir = parse_document(_make_pdf(), fmt="pdf", title="doc.pdf")
    assert not ir.meta["low_structure"]
    headings = [s.heading for _, s in ir.iter_sections() if s.heading]
    assert "Section 1" in headings and "Section 3" in headings
    all_text = " ".join(b.text for _, s in ir.iter_sections() for b in s.blocks)
    assert "Confidential" not in all_text  # колонтитул вырезан


def test_pdf_table_extracted_as_table_block() -> None:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 60), "Report", fontsize=18)
    page.insert_text((72, 100), "Quarterly numbers below.", fontsize=11)
    # Таблица с явными линиями — find_tables детектит её по сетке
    x0, y0, col_w, row_h = 72, 140, 150, 24
    cells = [["Region", "Revenue"], ["North", "100"], ["South", "250"]]
    for row_index, row in enumerate(cells):
        for col_index, value in enumerate(row):
            page.insert_text(
                (x0 + col_index * col_w + 4, y0 + row_index * row_h + 16), value, fontsize=10
            )
    for row_index in range(len(cells) + 1):
        y = y0 + row_index * row_h
        page.draw_line((x0, y), (x0 + 2 * col_w, y))
    for col_index in range(3):
        x = x0 + col_index * col_w
        page.draw_line((x, y0), (x, y0 + len(cells) * row_h))
    data: bytes = document.tobytes()
    document.close()

    ir = parse_document(data, fmt="pdf", title="report.pdf")
    blocks = [b for _, s in ir.iter_sections() for b in s.blocks]
    tables = [b for b in blocks if b.type == BlockType.TABLE]
    assert tables, "таблица не извлечена из PDF"
    assert "Revenue" in tables[0].text and "250" in tables[0].text
    # Текст таблицы не задублирован в параграфах
    paragraphs = " ".join(b.text for b in blocks if b.type == BlockType.PARAGRAPH)
    assert "250" not in paragraphs


def test_pdf_without_text_raises() -> None:
    document = pymupdf.open()
    document.new_page()
    empty: bytes = document.tobytes()
    document.close()
    with pytest.raises(ParserError):
        parse_document(empty, fmt="pdf", title="empty.pdf")


# --- confluence html ---

CONFLUENCE_HTML = """
<h1>Регламент</h1>
<p>Общее описание процесса.</p>
<h2>Шаги</h2>
<ul><li>Первый</li><li>Второй</li></ul>
<table><tr><th>Роль</th><th>Право</th></tr><tr><td>admin</td><td>всё</td></tr></table>
<ac:structured-macro ac:name="code">
  <ac:parameter ac:name="language">bash</ac:parameter>
  <ac:plain-text-body>make deploy</ac:plain-text-body>
</ac:structured-macro>
"""


def test_confluence_html_structure() -> None:
    ir = parse_confluence_html(CONFLUENCE_HTML, title="Регламент")
    paths = [path for path, _ in ir.iter_sections()]
    assert ["Регламент", "Шаги"] in paths
    blocks = [b for _, s in ir.iter_sections() for b in s.blocks]
    assert any(b.type == BlockType.TABLE and "Роль" in b.text for b in blocks)
    assert any(b.type == BlockType.CODE and "make deploy" in b.text for b in blocks)
    assert any(b.type == BlockType.LIST and "Первый" in b.text for b in blocks)
